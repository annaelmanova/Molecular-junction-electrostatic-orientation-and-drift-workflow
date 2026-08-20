from pathlib import Path
import json, sys
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

run=Path(sys.argv[1]).resolve(); out=run/'outputs'; figs=out/'figures'
cfg=json.loads((run/'user_inputs.json').read_text(encoding='utf-8'))
molecule=cfg.get('molecule_name','Y6').strip()
left_layer=cfg.get('left_layer_name','ZnO').strip()
interlayer=cfg.get('interlayer_name','SAM').strip()
active_layer=cfg.get('active_layer_name','PM6:Y6').strip()
summary=pd.read_csv(out/'data'/'final_drift_summary.csv')
doc=Document(); sec=doc.sections[0]
sec.top_margin=sec.bottom_margin=Inches(.75); sec.left_margin=sec.right_margin=Inches(.8)
for name,size,color in [('Normal',10,None),('Title',20,RGBColor(23,54,93)),('Heading 1',15,RGBColor(23,54,93)),('Heading 2',12,RGBColor(45,91,128)),('Caption',8.5,RGBColor(85,95,105))]:
 s=doc.styles[name]; s.font.name='Arial'; s.font.size=Pt(size)
 if color: s.font.color.rgb=color
def shade(c,fill):
 x=OxmlElement('w:shd'); x.set(qn('w:fill'),fill); c._tc.get_or_add_tcPr().append(x)
def table(rows,headers):
 t=doc.add_table(rows=1,cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
 for i,h in enumerate(headers):
  t.rows[0].cells[i].text=str(h); shade(t.rows[0].cells[i],'315B7A')
  for r in t.rows[0].cells[i].paragraphs[0].runs: r.font.color.rgb=RGBColor(255,255,255); r.bold=True
 for j,row in enumerate(rows):
  cs=t.add_row().cells
  for i,v in enumerate(row):
   cs[i].text=str(v)
   for p in cs[i].paragraphs:
    p.paragraph_format.space_after=Pt(0)
    for r in p.runs: r.font.name='Arial'; r.font.size=Pt(7.5)
  if j%2:
   for c in cs: shade(c,'F2F4F7')
 return t
def picture(path,caption,width=6.05):
 p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True
 p.add_run().add_picture(str(path),width=Inches(width))
 c=doc.add_paragraph(caption,style='Caption'); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
def equation(text):
 p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
 p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(6)
 r=p.add_run(text); r.italic=True; r.font.name='Cambria Math'; r.font.size=Pt(10.5)

doc.add_heading(f'{left_layer}/{interlayer}/{active_layer} junction with {molecule}: electrostatic, orientation, force and drift model',0)
doc.add_paragraph(f'Configured single-notebook calculation; run {run.name}.')
doc.add_heading('1. Workflow and scope',1)
doc.add_paragraph('The workflow propagates user-supplied junction and quantum-chemical inputs through electrostatics, analytical field derivatives, an isotropic molecular baseline, orientation-energy minimization, and anisotropic force and drift. The isotropic stage deliberately precedes tensor minimization so that electrostatic-profile effects can be separated from molecular anisotropy.')
p=figs/'Figure_01_workflow_or_cover.png'
if p.exists(): picture(p,'Figure 1. Workflow from device and molecular inputs to electrostatics, force components, orientation, and drift.',3.15)
doc.add_paragraph('All voltage, interlayer-step, screening-length, diffusion-coefficient, temperature, and time sweeps are carried through the same reproducible notebook. Drift lengths are local directed Einstein-relation displacement scales, not complete morphology trajectories.')

doc.add_heading('2. Junction electrostatics and model choice',1)
p=figs/'Figure_02_band_bending_scheme.png'
if p.exists(): picture(p,'Figure 2. Junction formation, Fermi-level equilibration, and interfacial band bending.',6.0)
doc.add_paragraph(f'Contact equilibration creates a spatially varying potential across the {left_layer}/{interlayer}/{active_layer} junction. Two limiting closures are compared. Exponential screening describes a continuously decaying interfacial perturbation; the parabolic depletion approximation corresponds to an approximately uniform space-charge density over a finite width.')
doc.add_heading('2.1 Exponential screening',2)
equation('φ_exp(z) = -(V_bi - V_app)z/d + Δφ_ZnO exp(-z/λ_ZnO) + Δφ_SAM exp(-z/λ_SAM)')
equation('E_exp(z) = -dφ/dz = (V_bi - V_app)/d + (Δφ_ZnO/λ_ZnO)e^(-z/λ_ZnO) + (Δφ_SAM/λ_SAM)e^(-z/λ_SAM)')
equation("E′_exp(z) = -(Δφ_ZnO/λ_ZnO²)e^(-z/λ_ZnO) -(Δφ_SAM/λ_SAM²)e^(-z/λ_SAM)")
equation("E″_exp(z) = (Δφ_ZnO/λ_ZnO³)e^(-z/λ_ZnO) +(Δφ_SAM/λ_SAM³)e^(-z/λ_SAM)")
doc.add_heading('2.2 Parabolic depletion reference',2)
equation('p(z;λ) = (1 - z/λ)² for 0 ≤ z < λ;    p(z;λ) = 0 for z ≥ λ')
equation('φ_par(z) = -(V_bi - V_app)z/d + Δφ_ZnO p(z;λ_ZnO) + Δφ_SAM p(z;λ_SAM)')
doc.add_paragraph('A uniform-field capacitor model has no interfacial gradient and therefore removes the leading permanent-dipole gradient force. An unscreened Coulomb model is singular at the interface and does not represent a distributed junction. A self-consistent Poisson or drift-diffusion solution would require dielectric profiles, mobile and trapped charge densities, injection conditions, and electrode boundary conditions that are not supplied. The exponential and parabolic forms are consequently transparent limiting models rather than claims of a unique microscopic charge distribution.')
fig_no=2
for name,cap in [('01_potential_and_field.png','Exponential versus parabolic potential and field for representative sweep groups.'),('01b_voltage_lambda_sweeps.png',f'Exponential {left_layer} decay-length sweep.'),('01c_voltage_lambda_SAM_sweeps.png',f'Exponential {interlayer} decay-length sweep.'),('10_fixed_case_model_components.png','Exponential and parabolic potential/field components for the configured fixed diagnostic case.'),('02_field_first_second_derivatives.png','Electric field and its first two analytical spatial derivatives for both closures.')]:
 p=figs/name
 if p.exists(): fig_no+=1; picture(p,f'Figure {fig_no}. {cap}')
doc.add_paragraph('The potential remains continuous. Exponential screening produces smooth derivatives, whereas the ideal parabolic closure gives a linearly varying field and constant field gradient inside its depletion width. Its abrupt depletion edge formally carries a distribution-valued curvature term; the plotted ordinary second derivative excludes that idealized edge impulse.')

doc.add_heading('3. Isotropic molecular model',1)
doc.add_paragraph('The isotropic reference suppresses orientation while retaining the molecular response scales. It uses the ground-state dipole magnitude and rotational invariants of the response tensors.')
equation('μ_iso = |μ|,    α_iso = Tr(α)/3,    Q_iso = Tr(Q)/3')
doc.add_heading('3.1 Dipolar, polarizability and quadrupolar forces',2)
equation("F_μ = μ_iso E′,    F_α = α_iso E E′,    F_Q = ⅙ Q_iso E″")
equation('F_total = F_μ + F_α + F_Q')
doc.add_paragraph('The permanent-dipole term is driven by the field gradient, the induced-polarization term by E·E′, and the quadrupolar term by field curvature. The signed components are summed before drift is evaluated.')
for name,cap in [('11_isotropic_analytical_derivatives.png','Isotropic analytical potential, field, first derivative, and second derivative for the fixed exponential case.'),('12_isotropic_fixed_force_components.png','Isotropic dipolar, polarizability, quadrupolar, and total force at V_app = -5 V and Δφ_SAM = -2 V.'),('13_isotropic_total_force_sweep.png','Isotropic total-force sweep for V_app = -1 and -5 V and Δφ_SAM = -2, -0.7, and +0.7 V.')]:
 p=figs/name
 if p.exists():
  fig_no+=1; picture(p,f'Figure {fig_no}. {cap}',2.56 if name in {'12_isotropic_fixed_force_components.png','13_isotropic_total_force_sweep.png'} else 6.05)
doc.add_heading('3.2 Isotropic drift-length sweeps',2)
equation('v_drift = D F_total/(k_B T),    L_drift = v_drift t')
doc.add_paragraph(f'Diffusion coefficient and elapsed time rescale the displacement, while applied voltage, the {interlayer} potential step, and screening length alter the electrostatic force. The atlases below retain all requested parameter values for the configured times.')
for tm in cfg.get('times_min',[10,60]):
 p=figs/f'06_Ldrift_atlas_{tm:g}min.png'
 if p.exists(): fig_no+=1; picture(p,f'Figure {fig_no}. Isotropic drift-length atlas at {tm:g} min for all requested D, V_app, and Δφ_SAM values.')

doc.add_heading('4. Orientation-energy minimization',1)
doc.add_paragraph('The anisotropic calculation rotates the ground-state dipole, polarizability tensor, quadrupole tensor, and XYZ geometry together in their common Cartesian frame. For a trial unit vector n along the device-field direction:')
equation('μ_eff = n·μ,    α_eff = nᵀαn,    Q_eff = nᵀQn')
equation("U(n,z) = -E(n·μ) - ½E²(nᵀαn) - ⅙E′(nᵀQn)")
equation('n*(z) = arg min U(n,z),    n ∈ S²')
doc.add_paragraph('The hard minimum is the zero-temperature orientation. The finite-temperature response uses normalized weights w(n,z) proportional to exp[-(U-U_min)/(k_B T_orientation)]. When symmetry-related minima exchange identity near E = 0, the angular branch can change abruptly although the minimum energy is continuous; this is a degeneracy, not a potential-energy peak.')
for name,cap in [('14_three_orientation_energy_minima.png','Three distinct minimum-energy orientation curves obtained by sweeping the full orientation sphere at every coordinate.'),('03_anisotropic_orientation_tensor.png','Exponential hard-minimum orientation angles and effective tensor projections along the junction coordinate.')]:
 p=figs/name
 if p.exists(): fig_no+=1; picture(p,f'Figure {fig_no}. {cap}',2.56 if name=='14_three_orientation_energy_minima.png' else 6.05)
doc.add_heading(f'4.1 Positioned {molecule} molecular views',2)
doc.add_paragraph(f'The unrotated {molecule} geometry and two selected minimum-energy orientations are shown in the common molecular/device Cartesian frame. The arrow denotes the ground-state dipole moment, not a transition dipole moment.')
p=figs/'14b_positioned_Y6_views.png'
if molecule.casefold()=='y6' and p.exists(): fig_no+=1; picture(p,f'Figure {fig_no}. Unrotated {molecule} and two positioned minimum-energy molecular views.',6.05)
elif molecule.casefold()!='y6': doc.add_paragraph(f'The packaged Y6 render is intentionally omitted for custom molecule {molecule}; supply molecule-specific views generated from its own XYZ geometry and tensors.')

doc.add_heading('5. Anisotropic force and drift',1)
equation("F_μ = μ_eff E′,    F_α = α_eff E E′,    F_Q = ⅙ Q_eff E″")
equation('F_total = F_μ + F_α + F_Q')
doc.add_paragraph('At every coordinate the force uses either the hard-minimum tensor projections or their finite-temperature orientational averages. The same Einstein relation as in the isotropic baseline then gives the anisotropic drift length.')
for name,cap in [('15_three_orientation_force_components.png','Dipolar, polarizability, quadrupolar, and total force for each of the three exponential minimum-energy orientations.'),('16_three_orientation_Ldrift.png','Drift length for the three exponential minimum-energy orientations.')]:
 p=figs/name
 if p.exists(): fig_no+=1; picture(p,f'Figure {fig_no}. {cap}',2.56)
for tm in cfg.get('times_min',[10,60]):
 p=figs/f'17_anisotropic_Ldrift_full_sweep_{tm:g}min.png'
 if p.exists(): fig_no+=1; picture(p,f'Figure {fig_no}. Exponential-only anisotropic hard-minimum drift atlas at {tm:g} min; panels are split so no panel exceeds eight curves.')

doc.add_heading('6. Interpretation and limitations',1)
doc.add_paragraph('A larger applied voltage need not produce a proportionally larger drift in the prescribed exponential model because V_app mainly changes the uniform part of E, whereas the permanent-dipole force depends on E′. Stronger voltage dependence would require voltage-dependent screening, charge injection, space charge, or self-consistent boundary conditions. The reported drift is a local transport scale and does not include packing, trapping, energetic disorder, concentration redistribution, or back-reaction on the field.')
doc.add_heading('Appendix A. User-input parameters',1)
table([[k,', '.join(map(str,v)) if isinstance(v,list) else v] for k,v in cfg.items()],['Parameter','Value'])
doc.add_page_break(); doc.add_heading('Appendix B. Detailed drift summary tables',1)
doc.add_paragraph('Values are in nanometres; the complete machine-readable table is final_drift_summary.csv.')
for D in cfg.get('diffusion_m2_s',[]):
 for tm in cfg.get('times_min',[]):
  mx=f'max_abs_L_nm_D{D:.0e}_t{tm:g}min'; rms=f'rms_L_nm_D{D:.0e}_t{tm:g}min'
  if mx not in summary or rms not in summary: continue
  doc.add_heading(f'D = {D:.0e} m² s⁻¹; t = {tm:g} min',2)
  rows=[[r['model'],f"{r['Vapp_V']:g}",f"{r['SAM_V']:+g}",f"{r[mx]:.5g}",f"{r[rms]:.5g}"] for _,r in summary.iterrows()]
  table(rows,['Model','Vapp (V)','SAM step (V)','Maximum |L| (nm)','RMS L (nm)'])
docx=out/'Molecular_Junction_Report.docx'; doc.save(docx); print(docx)
